# nvu/export.py
from io import BytesIO
from typing import Dict, Any, Optional
import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# -------------------- Köməkçilər --------------------
def _fmt_int(x: Optional[int]) -> str:
    if x is None:
        return "—"
    try:
        return f"{int(x):,}".replace(",", " ")
    except Exception:
        return str(x)

def _make_table_borderless(table):
    """Word cədvəlində sərhədləri söndür."""
    try:
        tbl = table._tbl
        tblPr = getattr(tbl, "tblPr", None)
        if tblPr is None and hasattr(tbl, "get_or_add_tblPr"):
            tblPr = tbl.get_or_add_tblPr()
        if tblPr is None:
            return
        NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
        for el in tblPr.xpath("./w:tblBorders", namespaces=NS):
            tblPr.remove(el)
    except Exception:
        pass

def _shade_cell(cell, fill_hex: str = "D9E1F2"):
    """Cədvəl hüceyrəsinə fon rəngi ver (header üçün)."""
    try:
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), fill_hex)
        tcPr.append(shd)
    except Exception:
        pass

def _set_table_cell_margins(table, top=80, bottom=80, left=80, right=80):
    """
    Cədvəldə bütün hüceyrələr üçün eyni margin (twips) – header və data sətrlərində simmetriya üçün.
    1 pt ≈ 20 twips, 80 twips ≈ 4 pt.
    """
    try:
        tbl = table._tbl
        tblPr = getattr(tbl, "tblPr", None)
        if tblPr is None and hasattr(tbl, "get_or_add_tblPr"):
            tblPr = tbl.get_or_add_tblPr()
        if tblPr is None:
            return
        cellMar = tblPr.find(qn("w:tblCellMar"))
        if cellMar is None:
            cellMar = OxmlElement("w:tblCellMar")
            tblPr.append(cellMar)

        def _set(side, val):
            el = cellMar.find(qn(f"w:{side}"))
            if el is None:
                el = OxmlElement(f"w:{side}")
                cellMar.append(el)
            el.set(qn("w:w"), str(val))
            el.set(qn("w:type"), "dxa")

        _set("top", top); _set("bottom", bottom); _set("left", left); _set("right", right)
    except Exception:
        pass

def _to_text(val) -> str:
    """NaN/boş/“nan” dəyərləri '—' kimi yaz, ədəd varsa '.0' at."""
    if pd.isna(val):
        return "—"
    s = str(val).strip()
    if s.lower() in ("nan", "none", ""):
        return "—"
    try:
        if isinstance(val, float) and float(val).is_integer():
            return _fmt_int(int(val))
        if isinstance(val, int):
            return _fmt_int(val)
    except Exception:
        pass
    return s

def _sanitize_df_for_docx(df: pd.DataFrame) -> pd.DataFrame:
    dfx = df.copy()
    for c in dfx.columns:
        dfx[c] = dfx[c].map(_to_text)
    return dfx

def _drop_blank_rows(df: pd.DataFrame, key_cols) -> pd.DataFrame:
    """Göstərilməsini istəmədiyimiz boş/Nan/“nan”/“—” sətirləri sil."""
    dfx = df.copy()
    mask = pd.Series(True, index=dfx.index)
    for c in key_cols:
        if c in dfx.columns:
            s = dfx[c].astype(str).str.strip()
            bad = s.isna() | (s == "") | s.str.lower().isin(["nan", "none", "—"])
            mask &= ~bad
    return dfx[mask].copy()

def _add_table(doc: Document, df: pd.DataFrame, add_rownum: bool = False) -> None:
    dfx = df.copy()
    if add_rownum and len(dfx) > 0:
        dfx.insert(0, "Sıra №", range(1, len(dfx) + 1))
    dfx = _sanitize_df_for_docx(dfx)

    table = doc.add_table(rows=1, cols=len(dfx.columns))
    table.allow_autofit = True
    _set_table_cell_margins(table, top=80, bottom=80, left=80, right=80)
    _make_table_borderless(table)

    # Header
    hdr = table.rows[0].cells
    for i, col in enumerate(dfx.columns):
        _shade_cell(hdr[i], "D9E1F2")
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(str(col))
        run.bold = True
        run.font.name = "Arial"
        run.font.size = Pt(11)

    # Rows
    for _, row in dfx.iterrows():
        cells = table.add_row().cells
        for j, col in enumerate(dfx.columns):
            p = cells[j].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(row[col])
            run.font.name = "Arial"
            run.font.size = Pt(11)

def _subset(df: pd.DataFrame, preferred_cols) -> pd.DataFrame:
    cols = [c for c in preferred_cols if c in df.columns]
    return df[cols].copy() if cols else df.copy()

# -------------------- DOCX --------------------
def export_docx(report: Dict[str, Any], source_filename: str = "") -> bytes:
    doc = Document()

    # Ümumi stil – Arial 12
    base = doc.styles["Normal"]
    base.font.name = "Arial"
    base.font.size = Pt(12)

    # Heading 1 – Arial 18, tünd mavi
    h1 = doc.styles["Heading 1"]
    h1.font.name = "Arial"
    h1.font.size = Pt(18)
    h1.font.color.rgb = RGBColor(0x12, 0x3A, 0x7A)

    # Heading 2 – Arial 14, mavi
    h2 = doc.styles["Heading 2"]
    h2.font.name = "Arial"
    h2.font.size = Pt(14)
    h2.font.color.rgb = RGBColor(0x1F, 0x5A, 0xB6)

    # Başlıq
    doc.add_heading("NVU Arayış Paneli — Hesabat", level=1)

    # Mənbə fayl əvəzinə hesabat tarixi (qırmızı)
    p = doc.add_paragraph()
    r1 = p.add_run("Hesabat tarixi: ")
    r1.bold = True
    r1.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
    p.add_run(pd.Timestamp.now().strftime(" %Y-%m-%d %H:%M"))

    # 1) Utilizatorlar
    doc.add_heading("1) Utilizatorlar üzrə qəbul edilən NV sayları — yekun", level=2)
    util = report.get("utilizator_counts", pd.DataFrame())
    if isinstance(util, pd.DataFrame) and not util.empty:
        util_out = util.copy()
        if util_out.shape[1] >= 2:
            util_out.iloc[:, 1] = pd.to_numeric(util_out.iloc[:, 1], errors="coerce").fillna(0).astype(int)
            total = int(util_out.iloc[:, 1].sum())
            util_out.loc[len(util_out), util_out.columns[0]] = "CƏM"
            util_out.loc[len(util_out) - 1, util_out.columns[1]] = total
        _add_table(doc, util_out, add_rownum=False)
    else:
        doc.add_paragraph("Məlumat yoxdur.")

    # 2) Təsnifat
    doc.add_heading("2) Təsnifatlar üzrə — yekun", level=2)
    calc = bool(report.get("tesnifat_settings", {}).get("calc_amounts", False))
    ready = report.get("tesnifat_table")
    if isinstance(ready, pd.DataFrame) and not ready.empty:
        pref = ["Kod", "Təsnifat", "Say"] + (["Cəmi (AZN)"] if calc else [])
        t = _subset(ready, pref)

        # Çıxışda göstəriləcək sütunları hazırla
        cols = []
        if "Kod" in t.columns:
            cols.append("Kod")
        elif "Təsnifat" in t.columns:
            cols.append("Təsnifat")
        if "Say" in t.columns:
            cols.append("Say")
        if calc and "Cəmi (AZN)" in t.columns:
            t["Cəmi (AZN)"] = pd.to_numeric(t["Cəmi (AZN)"], errors="coerce").fillna(0).astype(int)
            cols.append("Cəmi (AZN)")

        t_display = t[cols].copy()
        # 1-ci sütunu mütləq “Təsnifat” adı ilə göstər
        if len(t_display.columns) > 0:
            first_col = t_display.columns[0]
            t_display.rename(columns={first_col: "Təsnifat"}, inplace=True)

        _add_table(doc, t_display, add_rownum=False)

        if "Say" in t.columns:
            p = doc.add_paragraph()
            p.add_run(f"Cəm say: {_fmt_int(int(pd.to_numeric(t['Say'], errors='coerce').fillna(0).sum()))}").bold = True
        if calc and "Cəmi (AZN)" in t.columns:
            p = doc.add_paragraph()
            p.add_run(
                f"Ümumi məbləğ (AZN): {_fmt_int(int(pd.to_numeric(t['Cəmi (AZN)'], errors='coerce').fillna(0).sum()))}"
            ).bold = True
    else:
        base = report.get("tesnifat_counts", pd.DataFrame())
        t = _subset(base, ["Təsnifat", "Say"])
        _add_table(doc, t, add_rownum=False)

    # 3+) Digər bölmələr – dinamik başlıqlar + Sıra №
    meta = report.get("top_counts_meta", {})
    sections = [
        ("3) Təsdiqedici sənədin statusları — yekun", "tesdiq_status_totals",
         ["Təsdiq edici sənədin statusu", "Say"], False, ["Təsdiq edici sənədin statusu"]),
        ("4) Təhvil-təslim sənədinin statusları — yekun", "tehvil_status_totals",
         ["Təhvil-təslim sənədinin statusu", "Say"], False, ["Təhvil-təslim sənədinin statusu"]),
        (f"5) Top {meta.get('erizeci_N', 50)} Ərizəçi", "top_erizeci",
         ["Ərizəçinin tam adı", "Say"], True, None),
        (f"6) Marka Top {meta.get('marka_N', 20)}", "top_marka",
         ["Marka", "Say"], True, None),
        (f"7) Modellər üzrə Top {meta.get('model_N', 10)}", "top_model",
         ["Marka", "Model", "Say"], True, None),
        (f"8) Rəng Top {meta.get('reng_N', 10)}", "top_reng",
         ["Rəng", "Say"], True, None),
        ("9) NV yaşları 10illik intervallarda — yekun", "year_bins",
         ["Buraxılış ili", "Say"], False, None),
    ]
    for title, key, pref, add_no, drop_keys in sections:
        doc.add_heading(title, level=2)
        d = report.get(key, pd.DataFrame())
        if isinstance(d, pd.DataFrame) and not d.empty:
            dx = d.copy()

            # Status cədvəllərində boş/NaN dəyərləri TAMAMİLƏ çıxar
            if drop_keys:
                dx = _drop_blank_rows(dx, drop_keys)

            for col in dx.columns:
                if pd.api.types.is_numeric_dtype(dx[col]):
                    dx[col] = pd.to_numeric(dx[col], errors="coerce").fillna(0).astype(int)

            if not dx.empty:
                _add_table(doc, _subset(dx, pref), add_rownum=add_no)
            else:
                doc.add_paragraph("Məlumat yoxdur.")
        else:
            doc.add_paragraph("Məlumat yoxdur.")

    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()

# -------------------- XLSX --------------------
def export_xlsx(report: Dict[str, Any]) -> bytes:
    """
    XLSX export (openpyxl):
      • Mövcud bölmələri ayrıca vərəqlərə yazır (əgər boş deyilsə)
      • 'Parametrlər' vərəqinə top_counts_meta (əgər varsa)
      • 'PowerBI_Feed' vərəqinə report['powerbi_feed'] (əgər varsa)
      • Heç nə yazılmayıbsa 'README' vərəqi yaradır
    """
    from pandas import ExcelWriter
    from datetime import datetime

    bio = BytesIO()
    wrote_any = False

    with ExcelWriter(bio, engine="openpyxl") as writer:
        # 1) Adlandırılmış əsas vərəqlər
        named_sections = [
            ("utilizator_counts",   "Utilizatorlar"),
            ("tesnifat_table",      "Təsnifat"),
            ("tesnifat_counts",     "Təsnifat_baza"),
            ("tesdiq_status_totals","Təsdiq_statusu"),
            ("tehvil_status_totals","Təhvil_statusu"),
            ("top_erizeci",         "Top_Ərizəçi"),
            ("top_marka",           "Top_Marka"),
            ("top_model",           "Top_Model"),
            ("top_reng",            "Top_Rəng"),
            ("year_bins",           "İllər_üzrə"),
        ]
        for key, sheet in named_sections:
            df = report.get(key)
            if isinstance(df, pd.DataFrame) and not df.empty:
                df.to_excel(writer, sheet_name=sheet[:31], index=False)
                wrote_any = True

        # 2) Parametrlər (Top-N meta)
        meta = report.get("top_counts_meta")
        if isinstance(meta, dict) and meta:
            pd.DataFrame([meta]).to_excel(writer, sheet_name="Parametrlər", index=False)
            wrote_any = True

        # 3) Power BI feed (tek-sheet)
        feed = report.get("powerbi_feed")
        if isinstance(feed, pd.DataFrame) and not feed.empty:
            feed.to_excel(writer, sheet_name="PowerBI_Feed", index=False)
            wrote_any = True

        # 4) Ehtiyat README
        if not wrote_any:
            pd.DataFrame({
                "Info": ["No data tables were available to export."],
                "GeneratedAt": [datetime.now().strftime("%Y-%m-%d %H:%M:%S")],
                "Hint": ["Check filters and input file, then retry."],
            }).to_excel(writer, sheet_name="README", index=False)

    bio.seek(0)
    return bio.getvalue()
