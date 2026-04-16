import streamlit as st
import pandas as pd
import re
import runpy
from pathlib import Path
from io import BytesIO
from datetime import datetime
from typing import List

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

st.set_page_config(page_title="NVU Platform", layout="wide")

if "page" not in st.session_state:
    st.session_state.page = "home"

st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700&display=swap');
html, body, [class*="css"] {font-family: 'Inter', sans-serif;}
.stApp {background: linear-gradient(180deg, #f4f8fc 0%, #eaf2fb 100%);}
.main-title {text-align: center; font-size: 42px; font-weight: 700; color: #123b63; margin-top: 20px; margin-bottom: 40px;}
.card-wrap {background: #ffffff; border-radius: 22px; padding: 28px 22px; min-height: 220px; box-shadow: 0 10px 30px rgba(18,59,99,0.10); border: 1px solid #dbe7f3; display: flex; flex-direction: column; justify-content: center; align-items: center;}
.card-icon {font-size: 42px; margin-bottom: 18px;}
.card-title {text-align: center; font-size: 24px; font-weight: 700; color: #163a5c; line-height: 1.4;}
.card-sub {text-align: center; font-size: 15px; color: #5e7388; margin-top: 10px; line-height: 1.6;}
.stButton > button {width: 100%; border-radius: 12px; border: none; background: linear-gradient(90deg, #1f5f99 0%, #2d79bd 100%); color: white; font-size: 18px; font-weight: 600; padding: 12px 18px; margin-top: 14px;}
.result-box {background: white; border-radius: 16px; padding: 18px 20px; border: 1px solid #dbe7f3; box-shadow: 0 6px 18px rgba(18,59,99,0.08); margin-bottom: 12px; color: #163a5c; font-size: 17px; line-height: 1.7;}
.small-note {color: #5e7388; font-size: 14px; margin-top: -10px; margin-bottom: 20px;}
</style>
""", unsafe_allow_html=True)

# -------- 1-ci modul daxili menyu --------
def clean_label_from_filename(name: str) -> str:
    base = name.rsplit('.', 1)[0]
    base = re.sub(r'^\d+[_\- ]*', '', base)
    return base.replace('_', ' ').replace('📤', '').replace('📥', '').strip()

def module1_files():
    files = []
    if Path('0_NVU_Arayış_Paneli.py').exists():
        files.append(('Panel', '0_NVU_Arayış_Paneli.py'))
    mdir = Path('module1_pages')
    if mdir.exists():
        for p in sorted(mdir.glob('*.py')):
            files.append((clean_label_from_filename(p.name), str(p)))
    return files

def run_streamlit_page(path: str):
    runpy.run_path(path, run_name='__main__')

# -------- 2-ci modul --------
EXCEL_SHEET_DEFAULT = ''
PLACEHOLDER_OPTIONS = [
    'NETICELER VE SIYAHI BURA YAZILACAQ.',
    'NETICELER VE SIYAHI BURA YAZILACAQ',
    'NETICƏLƏR VƏ SİYAHI BURA YAZILACAQ.',
    'NETICƏLƏR VƏ SİYAHI BURA YAZILACAQ',
]
BOLD_LABEL_FOR_ALL = True
BOLD_NV = set()

def _norm(s: str) -> str:
    return re.sub(r'\s+', ' ', s.strip().lower())

def _has_placeholder_text(text: str) -> bool:
    t = _norm(text)
    return any(_norm(opt) in t for opt in PLACEHOLDER_OPTIONS)

def normalize_columns(df: pd.DataFrame) -> pd.DataFrame:
    def strip_diacritics(x: str) -> str:
        return (x.replace('ş','s').replace('Ş','S').replace('ı','i').replace('İ','I')
                 .replace('ə','e').replace('Ə','E').replace('ö','o').replace('Ö','O')
                 .replace('ü','u').replace('Ü','U').replace('ç','c').replace('Ç','C')
                 .replace('ğ','g').replace('Ğ','G'))
    col_satis = col_siyahi = None
    for c in df.columns:
        lc = strip_diacritics(str(c).lower().strip())
        if 'satis' in lc and ('siralama' in lc or 'siralamasi' in lc):
            col_satis = c
        if 'siyahi' in lc or 'siyah' in lc:
            col_siyahi = c
    if col_satis is None or col_siyahi is None:
        raise KeyError("Lazımi sütunlar tapılmadı. Gözlənilən: 'Satış sıralaması' və 'siyahı'.")
    return df.rename(columns={col_satis: 'Satis', col_siyahi: 'Nomre'})

def extract_numeric(series: pd.Series) -> pd.Series:
    return pd.to_numeric(series.astype(str).str.replace(r'\D', '', regex=True), errors='coerce')

def build_line_for_one_sale(df: pd.DataFrame, s: int) -> str:
    subset = df.loc[df['Satis'] == s, 'Num'].dropna().astype(int)
    nums = sorted(subset.unique().tolist())
    return f"{s}-ci NV: {', '.join(str(n) for n in nums)}"

def set_run_arial12(run, bold=False):
    run.font.name = 'Arial'
    if run._element.rPr is None:
        run._element.get_or_add_rPr()
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    run.font.size = Pt(12)
    run.bold = bool(bold)

def add_nv_line_to_paragraph(p, line: str, make_label_bold=True, bold_whole=False):
    m = re.match(r'^(\d+-ci NV:)(\s*)(.*)$', line)
    if not m:
        r = p.add_run(line)
        set_run_arial12(r, bold=bold_whole)
        return
    label, spaces, rest = m.groups()
    r1 = p.add_run(label)
    set_run_arial12(r1, bold=(bold_whole or make_label_bold))
    if spaces:
        r_sp = p.add_run(spaces)
        set_run_arial12(r_sp, bold=bold_whole and not make_label_bold)
    r2 = p.add_run(rest)
    set_run_arial12(r2, bold=bold_whole)

def collect_placeholders(doc: Document):
    found = []
    for p in doc.paragraphs:
        if _has_placeholder_text(''.join(r.text for r in p.runs)):
            found.append(p)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if _has_placeholder_text(''.join(r.text for r in p.runs)):
                        found.append(p)
    return found

def fill_placeholders(doc: Document, lines: List[str]):
    targets = collect_placeholders(doc)
    if not targets:
        raise FileNotFoundError('Word şablonunda placeholder tapılmadı.')
    p = targets[0]
    p.text = ''
    first = True
    for line in lines:
        if not first:
            p.add_run().add_break()
        first = False
        add_nv_line_to_paragraph(p, line, make_label_bold=BOLD_LABEL_FOR_ALL, bold_whole=False)

def build_output_name(sales_list: List[int]) -> str:
    stamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    return f"AKT_{stamp}__NV-{'-'.join(str(s) for s in sales_list)}.docx"

# -------- 3-cü modul --------
def to_amount(series: pd.Series) -> pd.Series:
    s = series.fillna('').astype(str)
    s = s.str.replace(' ', '', regex=False).str.replace(',', '.', regex=False)
    s = s.str.extract(r'([-+]?\d*\.?\d+)')[0]
    return pd.to_numeric(s, errors='coerce').fillna(0)

def process_amounts(df_raw: pd.DataFrame, sales_input: str):
    if df_raw.shape[1] < 5:
        return None, 'Faylda ən azı 5 sütun olmalıdır.'
    satis = pd.to_numeric(df_raw.iloc[:, 0], errors='coerce').ffill()
    mebleg = to_amount(df_raw.iloc[:, 4])
    df = pd.DataFrame({'Satis': satis, 'Mebleg': mebleg}).dropna(subset=['Satis']).copy()
    df['Satis'] = df['Satis'].astype(int)
    sales_list = [int(x) for x in re.split(r'[,\s;]+', sales_input.strip()) if x.isdigit()]
    if not sales_list:
        return None, 'Düzgün satış nömrəsi daxil edilmədi.'
    results, total_say, total_mebleg = [], 0, 0.0
    for s in sales_list:
        sub = df[df['Satis'] == s]
        say = int(sub['Mebleg'].gt(0).sum())
        meb = float(sub['Mebleg'].sum())
        total_say += say
        total_mebleg += meb
        results.append(f'{s}-ci NV: Say={say}, Məbləğ={meb:g}')
    results.append(f'TOTAL: Say={total_say}, Məbləğ={total_mebleg:g}')
    return results, None

# -------- Ana səhifə --------
if st.session_state.page == 'home':
    st.markdown('<div class="main-title">NVU Rəqəmsal İdarəetmə Platforması</div>', unsafe_allow_html=True)
    col1, col2, col3 = st.columns(3, gap='large')
    with col1:
        st.markdown('<div class="card-wrap"><div class="card-icon">📄</div><div class="card-title">Arayışların Avtomatlaşdırılmış Hazırlanması</div><div class="card-sub">Arayışların operativ və vahid formatda hazırlanması üçün modul</div></div>', unsafe_allow_html=True)
        if st.button('Giriş', key='btn1'):
            st.session_state.page = 'arayis'
            st.rerun()
    with col2:
        st.markdown('<div class="card-wrap"><div class="card-icon">🧾</div><div class="card-title">Aktların Avtomatlaşdırılmış Tərtibi</div><div class="card-sub">Excel məlumatlarının Word şablonuna avtomatik yerləşdirilməsi üçün modul</div></div>', unsafe_allow_html=True)
        if st.button('Giriş', key='btn2'):
            st.session_state.page = 'akt'
            st.rerun()
    with col3:
        st.markdown('<div class="card-wrap"><div class="card-icon">💳</div><div class="card-title">Maliyyə Ödənişlərinə Nəzarət və Yoxlama Sistemi</div><div class="card-sub">Satış sıralamasına görə say və məbləğin hesablanması üçün modul</div></div>', unsafe_allow_html=True)
        if st.button('Giriş', key='btn3'):
            st.session_state.page = 'odenis'
            st.rerun()

elif st.session_state.page == 'arayis':
    st.markdown('<div class="main-title">Arayışların Avtomatlaşdırılmış Hazırlanması</div>', unsafe_allow_html=True)
    files = module1_files()
    if not files:
        st.error("1-ci modul faylları tapılmadı. `0_NVU_Arayış_Paneli.py` və `module1_pages/` lazımdır.")
    else:
        labels = [x[0] for x in files]
        paths = {label: path for label, path in files}
        if 'module1_subpage' not in st.session_state:
            st.session_state['module1_subpage'] = labels[0]
        chosen = st.selectbox('Bölmə seç', labels, index=labels.index(st.session_state['module1_subpage']) if st.session_state['module1_subpage'] in labels else 0)
        st.session_state['module1_subpage'] = chosen
        st.markdown(f'**Aktiv bölmə:** {chosen}')
        st.divider()
        run_streamlit_page(paths[chosen])
    if st.button('← Geri qayıt', key='back_from_arayis'):
        st.session_state.page = 'home'
        st.rerun()

elif st.session_state.page == 'akt':
    st.markdown('<div class="main-title">Aktların Avtomatlaşdırılmış Tərtibi</div>', unsafe_allow_html=True)
    colf1, colf2 = st.columns([1, 1])
    with colf1:
        excel_file = st.file_uploader('Excel cədvəli', type=['xlsx'], key='akt_excel')
    with colf2:
        docx_file = st.file_uploader('Word faylı', type=['docx'], key='akt_docx')
    sheet = st.text_input('Vərəq adı (boş saxla = 1-ci vərəq)', value=EXCEL_SHEET_DEFAULT)
    sales_raw = st.text_input('NV satış nömrələri (vergüllə)', placeholder='məs: 1,2,3')
    if st.button('AKT yarat və endir', key='akt_go'):
        if not excel_file or not docx_file:
            st.error('Həm Excel, həm də Word faylı əlavə edilməlidir.')
        else:
            try:
                sales_list = [int(x.strip()) for x in sales_raw.split(',') if x.strip()]
                if not sales_list:
                    st.error('NV satış nömrələri boşdur.')
                    st.stop()
                df = pd.read_excel(excel_file, sheet_name=sheet.strip() if sheet.strip() else 0, dtype=object, engine='openpyxl')
                df = normalize_columns(df)[['Satis', 'Nomre']].copy()
                df['Satis'] = pd.to_numeric(df['Satis'], errors='coerce').ffill().astype('Int64')
                df['Num'] = extract_numeric(df['Nomre'])
                lines = [build_line_for_one_sale(df, s) for s in sales_list]
                doc = Document(docx_file)
                fill_placeholders(doc, lines)
                out_name = build_output_name(sales_list)
                buf = BytesIO(); doc.save(buf); buf.seek(0)
                st.success(f'Hazırdır: {out_name}')
                st.download_button('Docx endir', data=buf.getvalue(), file_name=out_name, mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            except Exception as e:
                st.error(f'Xəta: {e}')
    if st.button('← Geri qayıt', key='back_from_akt'):
        st.session_state.page = 'home'
        st.rerun()

elif st.session_state.page == 'odenis':
    st.markdown('<div class="main-title">Maliyyə Ödənişlərinə Nəzarət və Yoxlama Sistemi</div>', unsafe_allow_html=True)
    uploaded_file = st.file_uploader('Excel faylı əlavə et', type=['xlsx', 'xls'], key='odenis_excel')
    st.markdown('<div class="small-note">Qeyd: proqram həmişə faylın 1-ci sheet-ni götürür.</div>', unsafe_allow_html=True)
    sales_input = st.text_input('Satış nömrələri', placeholder='Məs: 1188,1220')
    df_raw = None
    if uploaded_file is not None:
        try:
            df_raw = pd.read_excel(uploaded_file, sheet_name=0)
            st.success('Excel faylı uğurla oxundu.')
        except Exception as e:
            st.error(f'Fayl oxunmadı: {e}')
    if st.button('Hesabla', key='odenis_calc'):
        if uploaded_file is None:
            st.warning('Əvvəlcə Excel faylı əlavə et.')
        elif not sales_input.strip():
            st.warning('Satış nömrələrini daxil et.')
        else:
            results, err = process_amounts(df_raw, sales_input)
            if err:
                st.error(err)
            else:
                for line in results:
                    st.markdown(f'<div class="result-box">{line}</div>', unsafe_allow_html=True)
    if st.button('← Geri qayıt', key='back_from_odenis'):
        st.session_state.page = 'home'
        st.rerun()
